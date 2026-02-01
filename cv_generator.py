from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from typing import Dict, List, Any
import logging
import tempfile
import os
import json

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CVGenerator:
    def __init__(self, json_file_path: str = None):
        self.doc = Document()
        self.json_file_path = json_file_path
        logger.info("CV Generator initialized with ATS-friendly template")
    
    def generate_cv(self, json_file_path: str = None) -> str:
        """
        Main function: Creates a professional ATS-friendly CV document
        Reads data from JSON file
        Returns: Path to the generated CV file
        """
        file_path = json_file_path or self.json_file_path

        if not file_path:
            raise ValueError("No JSON file path provided. Either pass to generate_cv() or initialize with path.")

        # Load data from JSON file
        json_data = self._load_json_data(file_path)
        
        logger.info("Starting ATS-friendly CV generation for: %s", 
                   json_data.get('full_name', 'Unknown'))
        
        try:
            # Setup document styles for ATS compatibility
            self._setup_ats_document_styles()
            
            # Build CV sections in ATS-friendly order
            self._create_ats_header_section(json_data)
            self._create_professional_summary(json_data)
            self._create_work_experience(json_data)
            self._create_skills_section(json_data)
            self._create_education_section(json_data)
            self._create_languages_section(json_data)
            self._create_certifications_section(json_data)
            self._create_projects_section(json_data)
            
            # Save to temporary file
            cv_file_path = self._save_document()
            logger.info("ATS-friendly CV successfully generated at: %s", cv_file_path)
            
            return cv_file_path
            
        except Exception as e:
            logger.error("Error generating CV: %s", str(e))
            raise

    def _load_json_data(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            logger.info("Successfully loaded JSON data from: %s", file_path)
            logger.info("Data contains %d top-level fields", len(data))
            return data
        except FileNotFoundError:
            logger.error("JSON file not found: %s", file_path)
            raise
        except json.JSONDecodeError as e:
            logger.error("Invalid JSON format in file %s: %s", file_path, e)
            raise
        except Exception as e:
            logger.error("Error loading JSON file %s: %s", file_path, e)
            raise
    
    def _setup_ats_document_styles(self):
        """Setup ATS-friendly document styles"""
        # Set page margins 
        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(14)  
            section.page_width = Inches(8.5)   
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
    
    def _create_ats_header_section(self, data: Dict[str, Any]):
        """Create ATS-friendly header with name, title, and contact info"""
        
        # Name 
        name_paragraph = self.doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(data.get('full_name', '').upper())
        name_run.font.size = Pt(36)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        
        # Job Title - Centered below name (like "Digital Marketing | SEO | SEM")
        if data.get('job_title'):
            job_paragraph = self.doc.add_paragraph()
            job_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            job_run = job_paragraph.add_run(data.get('job_title', ''))
            job_run.font.size = Pt(14)
            job_run.font.name = 'Calibri'
            job_run.bold = True
        
        # Contact Information 
        contact_parts = []
        if data.get('city'):
            contact_parts.append(data['city'])
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('linkedin'):
            # Remove any "LinkedIn:" prefix if present, just show the URL
            linkedin = data['linkedin'].replace('LinkedIn:', '').strip()
            contact_parts.append(linkedin)
        if data.get('portfolio'):
            portfolio = data['portfolio'].replace("Portfolio:", "").strip()
            contact_parts.append(portfolio)
        
        if contact_parts:
            contact_paragraph = self.doc.add_paragraph()
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_paragraph.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(12)
            contact_run.font.name = 'Calibri'
        
    def _create_professional_summary(self, data: Dict[str, Any]):
        """Create professional summary section (2nd section after header)"""
        summary = data.get('professional_summary')
        
        if summary:
            self._add_ats_section_title("PROFESSIONAL SUMMARY")
            
            # Professional summary paragraph (normal font, not bold)
            summary_paragraph = self.doc.add_paragraph()
            summary_run = summary_paragraph.add_run(summary)
            summary_run.font.name = 'Calibri'
            summary_run.font.size = Pt(11)
                      
        else:
            logger.warning("No professional summary provided")
    
    def _create_work_experience(self, data: Dict):
        """Create work experience section with bullet points"""
        work_experience = data.get('work_experience')
        
        if work_experience:
            self._add_ats_section_title("WORK EXPERIENCE")
            
            for job in work_experience:
                # Job title and company
                job_paragraph = self.doc.add_paragraph()
                job_run = job_paragraph.add_run(f"{job.get('position', '')} | {job.get('company', '')}")
                job_run.font.name = 'Calibri'
                job_run.font.size = Pt(12)
                job_run.font.bold = True
                
                # Date and location
                date_location = []
                if job.get('start_date'):
                    date_location.append(job['start_date'])
                if job.get('end_date'):
                    date_location.append(f" - {job['end_date']}")
                if job.get('location'):
                    date_location.append(f" | {job['location']}")
                
                if date_location:
                    date_paragraph = self.doc.add_paragraph()
                    date_run = date_paragraph.add_run(''.join(date_location))
                    date_run.font.name = 'Calibri'
                    date_run.font.size = Pt(12)
                    date_run.font.bold=True
                
                # Responsibilities
                responsibilities = job.get('responsibilities', [])
                for responsibility in responsibilities:
                    responsibility_paragraph = self.doc.add_paragraph()
                    responsibility_paragraph.style = 'List Bullet'
                    responsibility_run = responsibility_paragraph.add_run(responsibility)
                    responsibility_run.font.name = 'Calibri'
                    responsibility_run.font.size = Pt(11)
            
        else:
            logger.info("No work experience provided - section skipped")

    def _create_skills_section(self, data: Dict[str, Any]):
        """Create skills section combining technical and soft skills"""
        technical_skills = data.get('technical_skills', [])
        soft_skills = data.get('soft_skills', [])
        
        if technical_skills or soft_skills:
            self._add_ats_section_title("SKILLS")
            
            skills_paragraph = self.doc.add_paragraph()
            
            # Technical skills
            if technical_skills:
                tech_skills_text ="Technical Skills: " + ", ".join(technical_skills)
                run = skills_paragraph.add_run(tech_skills_text)
                run.font.name='Calibri'
                run.font.size = Pt(11)
                skills_paragraph.add_run("\n")  # New line
            
            # Soft skills
            if soft_skills:
                soft_skills_text ="Soft Skills: " + ", ".join(soft_skills)
                run = skills_paragraph.add_run(soft_skills_text)
                run.font.name='Calibri'
                run.font.size = Pt(11)
                        
        else:
            logger.warning("No skills information provided")
    
    def _create_education_section(self, data: Dict):
        """Create education section"""
        education = data.get('education',[])
        
        if education:
            self._add_ats_section_title("EDUCATION")
            
            for edu in education:
                # LINE 1: Degree | GPA
                degree = edu.get('degree', '')
                gpa = edu.get('gpa')
                line1_text = degree
                if gpa:
                    line1_text += f" | GPA: {gpa}"
                line1 = self.doc.add_paragraph()
                run1 = line1.add_run(line1_text)
                run1.font.name = 'Calibri'
                run1.font.size = Pt(12)
                run1.font.bold = True
                
                # LINE 2: institution | Location | Dates
                university = edu.get('institution', '')
                location = edu.get('location', '')
                start_date = edu.get('start_date', '')
                end_date = edu.get('end_date', '')

                dates = ""
                if start_date and end_date:
                    dates = f"{start_date} - {end_date}"
                elif start_date:
                    dates = f"{start_date}"

                parts = [p for p in [university, location, dates] if p]
                line2_text = " | ".join(parts)

                line2 = self.doc.add_paragraph()
                run2 = line2.add_run(line2_text)
                run2.font.name = 'Calibri'
                run2.font.size = Pt(12)
                            
        else:
            logger.warning("No education information provided")

    def _create_languages_section(self, data: Dict[str, Any]):
        """Create languages section from list of dictionaries"""
        languages = data.get('languages',[])
        
        if languages:
            self._add_ats_section_title("LANGUAGES")
            
            languages_text = []
            for lang in languages:
                language_name = lang.get('name', '')
                proficiency = lang.get('proficiency', '')
                if language_name and proficiency:
                    lang_paragraph = self.doc.add_paragraph()
                    lang_paragraph.style = 'List Bullet'
            
                    # Language name bold
                    run_name = lang_paragraph.add_run(f"{language_name}")
                    run_name.font.bold = True
                    run_name.font.name = 'Calibri'
                    run_name.font.size = Pt(11)

                    # Colon and proficiency
                    run_prof = lang_paragraph.add_run(f": {proficiency}")
                    run_prof.font.bold = False
                    run_prof.font.name = 'Calibri'
                    run_prof.font.size = Pt(11)
            
        else:
            logger.info("No languages provided - section skipped")
    
    def _create_certifications_section(self, data: Dict[str, Any]):
        """Create certifications section"""
        certificates = data.get('certificates',[])
        
        if certificates:
            self._add_ats_section_title("CERTIFICATIONS")
            for cert in certificates:
                cert_paragraph = self.doc.add_paragraph()
                cert_paragraph.style = 'List Bullet'
                cert_run = cert_paragraph.add_run(cert)
                cert_run.font.name = 'Calibri'
                cert_run.font.size = Pt(11)
        else:
            logger.info("No certifications provided - section skipped")

    def _create_projects_section(self, data: Dict[str, Any]):
        """Create projects section from list of dictionaries"""
        projects = data.get('projects', [])
        
        if projects:
            self._add_ats_section_title("PROJECTS")
            
            for project in projects:
                # Project name
                project_paragraph = self.doc.add_paragraph()
                project_run = project_paragraph.add_run(project.get('name', ''))
                project_run.font.name = 'Calibri'
                project_run.font.size = Pt(12)
                project_run.font.bold = True
                
                # Project description
                if project.get('description'):
                    desc_paragraph = self.doc.add_paragraph()
                    desc_run = desc_paragraph.add_run(project['description'])
                    desc_run.font.name = 'Calibri'
                    desc_run.font.size = Pt(11)
                
                '''# Technologies used
                if project.get('technologies'):
                    tech_paragraph = self.doc.add_paragraph()
                    tech_run = tech_paragraph.add_run(f"Technologies: {', '.join(project['technologies'])}")
                    tech_run.font.name = 'Calibri'
                    tech_run.font.size = Pt(11)'''

                # Check for 'Project_Link' (Prompt instruction) AND 'project_link' (Common AI output)
                links = project.get('Project_Link') or project.get('project_link') or []
                
                # Ensure links is a list (sometimes AI returns a string by mistake)
                if isinstance(links, str):
                    links = [links]
                if links is None:
                    links = []

                if links:
                    tech_paragraph = self.doc.add_paragraph()
                    # Changed label from "Project_Link:" to "Link:" for better looks
                    tech_run = tech_paragraph.add_run(f"Link: {', '.join(links)}")
                    tech_run.font.name = 'Calibri'
                    tech_run.font.size = Pt(11)
                           
        else:
            logger.info("No projects provided - section skipped")
    
    def _add_ats_section_title(self, title: str):
        """
        Add ATS-friendly section title (all caps, bold, with proper spacing)
        Exactly like in the template: PROFESSIONAL SUMMARY, WORK EXPERIENCE, etc.
        """
        title_paragraph = self.doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(13)
        title_run.font.bold = True
        title_run.font.name = 'Calibri'
        title_run.font.all_caps = True
        
        # Add proper spacing after title 
        title_paragraph.paragraph_format.space_after = Pt(0)

        # Add horizontal line across the page width
        self._add_horizontal_line()

    def _add_horizontal_line(self):
        """Add horizontal line using text characters"""
        line_paragraph = self.doc.add_paragraph()
        line_run = line_paragraph.add_run("_" * 145)  # Using em-dash for solid line
        line_run.font.size = Pt(7)
        line_run.font.name = 'Calibri'
        line_run.font.color.rgb = None  # Black color
        line_paragraph.paragraph_format.space_after = Pt(6)
        line_paragraph.paragraph_format.space_before = Pt(0)
        line_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align left
    
    def _save_document(self) -> str:
        """Save the document to a temporary file and return the path"""
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        self.doc.save(temp_file.name)
        logger.info("Document saved to temporary file: %s", temp_file.name)
        return temp_file.name

if __name__ == "__main__":
    """
    Example usage when running this file directly
    """
    print("CV Generator - Direct JSON File Test")
    print("=" * 50)

    print("\nPass JSON file path to generate_cv()")
    generator2 = CVGenerator()
    cv_path2 = generator2.generate_cv("Temporary_File_2.json")
    print(f"CV generated at: {cv_path2}")

